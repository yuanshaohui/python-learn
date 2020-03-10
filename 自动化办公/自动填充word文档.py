from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  #自然段位置
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  #中文格式
#以上为docx库中用到的部分
import time


#文档中变化的内容
price = input("请输入今日的价格：")
company_list = ["客户一", "客户二", "客户三"]
today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y ="年", m = "月", d = "日")


#循环制作
for i in company_list:
    document = Document()  #新建文档
    document.styles["Normal"].font.name = u"宋体"  #文档样式
    #新建自然段
    p1 = document.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  #居中
    #设置第一行,内容，字体，字号，是否加粗
    run1 = p1.add_run("关于下达%s产品的价格通知"%(today))
    run1.font.name = "微软雅黑"
    run1.font.size = Pt(21)
    run1.font.bold = True
    #段落设置(段前段后各五榜)
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)
    document.save("%s-价格通知.docx" % i)

